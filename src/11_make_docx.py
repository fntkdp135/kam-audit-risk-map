# -*- coding: utf-8 -*-
"""⑪ 마크다운 원고 → 워드(.docx) 변환. 맑은 고딕, 표 지원."""
import os, re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")
DOCDIR = r"C:\Users\samsung-user\Desktop\삼일감사포트폴리오\자소서_기술노트"
FONT = "맑은 고딕"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)


def setfont(run, size=10, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_runs(par, text, base=10, color=None):
    """**굵게**, `코드` 인라인 처리."""
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            setfont(par.add_run(tok[2:-2]), base, bold=True, color=ACCENT)
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(base - 1)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.color.rgb = GREY
        else:
            setfont(par.add_run(tok), base, color=color)


def md_table(doc, rows):
    cols = len(rows[0])
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for c, txt in zip(cells, row):
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_runs(p, txt.strip(), base=9)
            if i == 0:
                for r in p.runs:
                    r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def convert(md_path, docx_path, title):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.2)
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.32

    i, buf_tbl = 0, []
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()

        # 표
        if s.startswith("|") and s.endswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if not re.match(r"^[\s\-:|]+$", s):
                buf_tbl.append(cells)
            i += 1
            nxt = lines[i].strip() if i < len(lines) else ""
            if not (nxt.startswith("|") and nxt.endswith("|")):
                if buf_tbl:
                    md_table(doc, buf_tbl)
                buf_tbl = []
            continue

        # 코드블록
        if s.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"
            r.font.size = Pt(8.8)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.color.rgb = GREY
            continue

        if not s:
            i += 1
            continue
        if s.startswith("---"):
            i += 1
            continue

        # 제목
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            lvl, txt = len(m.group(1)), m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(13 if lvl <= 2 else 8)
            p.paragraph_format.space_after = Pt(5)
            setfont(p.add_run(txt), {1: 15, 2: 12.5, 3: 11, 4: 10.5}[lvl],
                    bold=True, color=ACCENT if lvl <= 2 else None)
            i += 1
            continue

        # 인용(예상 질문)
        if s.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip("> ").rstrip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            add_runs(p, " ".join(block), base=10, color=GREY)
            continue

        # 목록
        m = re.match(r"^([-*])\s+(.*)$", s)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, m.group(2))
            i += 1
            continue
        m = re.match(r"^(\d+)\.\s+(.*)$", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, m.group(2))
            i += 1
            continue

        # 본문 (다음 빈 줄까지 이어붙임)
        para = [s]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
                r"^(#{1,4}\s|[-*]\s|\d+\.\s|>|\||```|---)", lines[i].strip()):
            para.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, " ".join(para))

    doc.save(docx_path)
    print(f"저장: {os.path.basename(docx_path)}")


if __name__ == "__main__":
    for md, out, ttl in [
        ("04_KAM감사위험지도_자소서초안.md", "04_KAM감사위험지도_자소서초안.docx", "자소서 초안"),
        ("04_KAM감사위험지도_기술설명_인터뷰대비.md",
         "04_KAM감사위험지도_기술설명_인터뷰대비.docx", "기술노트"),
    ]:
        convert(os.path.join(DOCDIR, md), os.path.join(DOCDIR, out), ttl)
